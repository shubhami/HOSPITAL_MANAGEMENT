from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
document = Document()
document.add_heading('This is the title', 0)
p = document.add_paragraph('And this is text ')
p.add_run('some bold text').bold = True
p.add_run('and italic text.').italic = True
document.save('output.docx')
